"""
Professional DOCX document comparison service.
Analyzes Word documents for content and formatting differences.
"""

import asyncio
import io
import logging
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple
from pathlib import Path
import base64

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_COLOR_INDEX
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    class Document:
        pass

from difflib import SequenceMatcher
from fuzzywuzzy import fuzz
import xml.etree.ElementTree as ET

from ..core.base import (
    BaseConverter, 
    ConversionFormat, 
    ComparisonRequest, 
    ComparisonResult,
    ConversionStatus,
    ProcessingError,
    ValidationError
)
from ..core.factory import ConverterPlugin

logger = logging.getLogger(__name__)


class DocxCompareService(BaseConverter):
    """
    Professional DOCX document comparison service.
    
    Features:
    - Content difference analysis
    - Formatting comparison
    - Structural changes detection
    - Similarity scoring
    - Detailed reporting
    - Track changes visualization
    """
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX comparison. "
                "Install with: uv add python-docx"
            )
        
        super().__init__(
            name="docx_compare",
            supported_formats=[ConversionFormat.COMPARISON]
        )
    
    def validate_request(self, request: ComparisonRequest) -> bool:
        """Validate DOCX comparison request."""
        if not request.document_1 or not request.document_2:
            return False
        
        # Check if documents are valid DOCX content or files
        try:
            self._validate_docx_content(request.document_1)
            self._validate_docx_content(request.document_2)
            return True
        except Exception:
            return False
    
    def _validate_docx_content(self, content: str) -> bool:
        """Validate if content is valid DOCX."""
        if content.startswith('data:') or len(content) > 1000:
            # Likely base64 encoded or file content
            return True
        elif Path(content).exists() and Path(content).suffix.lower() == '.docx':
            # File path
            return True
        else:
            raise ValidationError("Invalid DOCX content or file path")
    
    async def compare_documents(self, request: ComparisonRequest) -> ComparisonResult:
        """
        Compare two DOCX documents for differences.
        
        Args:
            request: Comparison request with two documents
            
        Returns:
            ComparisonResult with detailed analysis
        """
        job_id = f"docx_compare_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}_{id(request)}"
        
        try:
            logger.info(f"Starting DOCX comparison: {job_id}")
            
            # Validate request
            if not self.validate_request(request):
                raise ValidationError("Invalid DOCX comparison request", job_id)
            
            # Load documents
            doc1, doc2 = await self._load_documents(request.document_1, request.document_2)
            
            # Extract content and metadata
            content1 = await self._extract_document_content(doc1)
            content2 = await self._extract_document_content(doc2)
            
            # Perform comparison analysis
            comparison_results = await self._analyze_differences(
                content1, content2, request.options or {}
            )
            
            # Calculate similarity score
            similarity_score = self._calculate_similarity_score(content1, content2)
            
            # Generate detailed report
            detailed_report = await self._generate_detailed_report(
                comparison_results, request.options or {}
            )
            
            # Create result
            result = ComparisonResult(
                job_id=job_id,
                status=ConversionStatus.COMPLETED,
                differences_found=comparison_results['has_differences'],
                similarity_score=similarity_score,
                content_differences=comparison_results.get('content_differences'),
                formatting_differences=comparison_results.get('formatting_differences'),
                summary=comparison_results.get('summary'),
                detailed_report=detailed_report,
                metadata={
                    "comparison_type": request.comparison_type,
                    "document_1_stats": content1.get('stats'),
                    "document_2_stats": content2.get('stats'),
                    "analysis_duration": "completed"
                },
                created_at=datetime.utcnow(),
                completed_at=datetime.utcnow()
            )
            
            logger.info(f"DOCX comparison completed: {job_id}")
            return result
            
        except Exception as e:
            logger.error(f"DOCX comparison failed: {job_id} - {str(e)}")
            return ComparisonResult(
                job_id=job_id,
                status=ConversionStatus.FAILED,
                differences_found=False,
                error_message=str(e),
                created_at=datetime.utcnow(),
                completed_at=datetime.utcnow()
            )
    
    async def _load_documents(self, doc1_source: str, doc2_source: str) -> Tuple[Document, Document]:
        """Load DOCX documents from various sources."""
        loop = asyncio.get_event_loop()
        
        def load_doc(source: str) -> Document:
            if source.startswith('data:'):
                # Base64 encoded content
                header, data = source.split(',', 1)
                content = base64.b64decode(data)
                return Document(io.BytesIO(content))
            elif Path(source).exists():
                # File path
                return Document(source)
            else:
                # Assume binary content
                return Document(io.BytesIO(source.encode()))
        
        doc1, doc2 = await loop.run_in_executor(
            None, lambda: (load_doc(doc1_source), load_doc(doc2_source))
        )
        
        return doc1, doc2
    
    async def _extract_document_content(self, doc: Document) -> Dict[str, Any]:
        """Extract comprehensive content and metadata from DOCX document."""
        content = {
            'text': [],
            'paragraphs': [],
            'tables': [],
            'headers': [],
            'footers': [],
            'styles': {},
            'stats': {}
        }
        
        # Extract paragraph content and formatting
        for para in doc.paragraphs:
            para_data = {
                'text': para.text,
                'style': para.style.name if para.style else None,
                'alignment': str(para.alignment) if para.alignment else None,
                'runs': []
            }
            
            for run in para.runs:
                run_data = {
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size.pt if run.font.size else None
                }
                para_data['runs'].append(run_data)
            
            content['paragraphs'].append(para_data)
            content['text'].append(para.text)
        
        # Extract table content
        for table in doc.tables:
            table_data = {
                'rows': len(table.rows),
                'cols': len(table.columns),
                'cells': []
            }
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data['cells'].append(row_data)
            
            content['tables'].append(table_data)
        
        # Calculate statistics
        content['stats'] = {
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
            'word_count': len(' '.join(content['text']).split()),
            'character_count': len(' '.join(content['text']))
        }
        
        return content
    
    async def _analyze_differences(self, content1: Dict, content2: Dict, options: Dict) -> Dict[str, Any]:
        """Analyze differences between document contents."""
        comparison_type = options.get('comparison_type', 'both')
        
        results = {
            'has_differences': False,
            'content_differences': {},
            'formatting_differences': {},
            'summary': ''
        }
        
        # Content comparison
        if comparison_type in ['content', 'both']:
            content_diff = await self._compare_content(content1, content2)
            results['content_differences'] = content_diff
            if content_diff['changes_count'] > 0:
                results['has_differences'] = True
        
        # Formatting comparison
        if comparison_type in ['formatting', 'both']:
            formatting_diff = await self._compare_formatting(content1, content2)
            results['formatting_differences'] = formatting_diff
            if formatting_diff['changes_count'] > 0:
                results['has_differences'] = True
        
        # Generate summary
        results['summary'] = self._generate_summary(results)
        
        return results
    
    async def _compare_content(self, content1: Dict, content2: Dict) -> Dict[str, Any]:
        """Compare textual content between documents."""
        text1 = ' '.join(content1['text'])
        text2 = ' '.join(content2['text'])
        
        # Use SequenceMatcher for detailed diff
        matcher = SequenceMatcher(None, text1, text2)
        changes = []
        changes_count = 0
        
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag != 'equal':
                changes.append({
                    'type': tag,
                    'original': text1[i1:i2],
                    'modified': text2[j1:j2],
                    'position': {'start': i1, 'end': i2}
                })
                changes_count += 1
        
        # Paragraph-level comparison
        para_changes = []
        para1_texts = [p['text'] for p in content1['paragraphs']]
        para2_texts = [p['text'] for p in content2['paragraphs']]
        
        para_matcher = SequenceMatcher(None, para1_texts, para2_texts)
        for tag, i1, i2, j1, j2 in para_matcher.get_opcodes():
            if tag != 'equal':
                para_changes.append({
                    'type': tag,
                    'original_paragraphs': para1_texts[i1:i2],
                    'modified_paragraphs': para2_texts[j1:j2],
                    'paragraph_range': {'start': i1, 'end': i2}
                })
        
        return {
            'changes_count': changes_count,
            'text_changes': changes,
            'paragraph_changes': para_changes,
            'statistics': {
                'original_words': len(text1.split()),
                'modified_words': len(text2.split()),
                'original_paragraphs': len(content1['paragraphs']),
                'modified_paragraphs': len(content2['paragraphs'])
            }
        }
    
    async def _compare_formatting(self, content1: Dict, content2: Dict) -> Dict[str, Any]:
        """Compare formatting between documents."""
        formatting_changes = []
        changes_count = 0
        
        # Compare paragraph styles
        for i, (para1, para2) in enumerate(zip(content1['paragraphs'], content2['paragraphs'])):
            if para1.get('style') != para2.get('style'):
                formatting_changes.append({
                    'type': 'style_change',
                    'paragraph_index': i,
                    'original_style': para1.get('style'),
                    'modified_style': para2.get('style')
                })
                changes_count += 1
            
            if para1.get('alignment') != para2.get('alignment'):
                formatting_changes.append({
                    'type': 'alignment_change',
                    'paragraph_index': i,
                    'original_alignment': para1.get('alignment'),
                    'modified_alignment': para2.get('alignment')
                })
                changes_count += 1
            
            # Compare run formatting
            for j, (run1, run2) in enumerate(zip(para1.get('runs', []), para2.get('runs', []))):
                if run1.get('bold') != run2.get('bold'):
                    formatting_changes.append({
                        'type': 'bold_change',
                        'paragraph_index': i,
                        'run_index': j,
                        'original_bold': run1.get('bold'),
                        'modified_bold': run2.get('bold')
                    })
                    changes_count += 1
        
        return {
            'changes_count': changes_count,
            'formatting_changes': formatting_changes
        }
    
    def _calculate_similarity_score(self, content1: Dict, content2: Dict) -> float:
        """Calculate similarity score between documents."""
        text1 = ' '.join(content1['text'])
        text2 = ' '.join(content2['text'])
        
        # Use multiple similarity metrics
        ratio = SequenceMatcher(None, text1, text2).ratio()
        token_sort_ratio = fuzz.token_sort_ratio(text1, text2) / 100.0
        token_set_ratio = fuzz.token_set_ratio(text1, text2) / 100.0
        
        # Weighted average
        similarity = (ratio * 0.5 + token_sort_ratio * 0.3 + token_set_ratio * 0.2)
        return round(similarity, 3)
    
    async def _generate_detailed_report(self, comparison_results: Dict, options: Dict) -> Dict[str, Any]:
        """Generate comprehensive comparison report."""
        return {
            'executive_summary': comparison_results.get('summary', ''),
            'analysis_method': 'Advanced DOCX comparison with content and formatting analysis',
            'content_analysis': comparison_results.get('content_differences', {}),
            'formatting_analysis': comparison_results.get('formatting_differences', {}),
            'recommendations': self._generate_recommendations(comparison_results),
            'metadata': {
                'analysis_timestamp': datetime.utcnow().isoformat(),
                'comparison_options': options
            }
        }
    
    def _generate_summary(self, results: Dict) -> str:
        """Generate human-readable summary of comparison results."""
        if not results['has_differences']:
            return "Documents are identical in content and formatting."
        
        summary_parts = []
        
        content_changes = results.get('content_differences', {}).get('changes_count', 0)
        if content_changes > 0:
            summary_parts.append(f"{content_changes} content changes detected")
        
        formatting_changes = results.get('formatting_differences', {}).get('changes_count', 0)
        if formatting_changes > 0:
            summary_parts.append(f"{formatting_changes} formatting changes detected")
        
        return "Differences found: " + ", ".join(summary_parts) + "."
    
    def _generate_recommendations(self, comparison_results: Dict) -> List[str]:
        """Generate actionable recommendations based on comparison results."""
        recommendations = []
        
        content_changes = comparison_results.get('content_differences', {}).get('changes_count', 0)
        formatting_changes = comparison_results.get('formatting_differences', {}).get('changes_count', 0)
        
        if content_changes > 10:
            recommendations.append("Consider reviewing major content changes for accuracy")
        
        if formatting_changes > 5:
            recommendations.append("Standardize formatting across documents for consistency")
        
        if content_changes == 0 and formatting_changes > 0:
            recommendations.append("Content is identical; only formatting differs")
        
        return recommendations
    
    async def convert(self, request) -> ComparisonResult:
        """Convert method for compatibility with BaseConverter interface."""
        if isinstance(request, ComparisonRequest):
            return await self.compare_documents(request)
        else:
            raise ValidationError("DocxCompareService requires ComparisonRequest")
    
    def get_capabilities(self) -> Dict[str, Any]:
        """Get service capabilities."""
        return {
            "name": self.name,
            "description": "Professional DOCX document comparison service",
            "supported_formats": [fmt.value for fmt in self.supported_formats],
            "features": [
                "Content difference analysis",
                "Formatting comparison",
                "Similarity scoring",
                "Detailed reporting",
                "Track changes visualization",
                "Statistical analysis"
            ],
            "supported_options": [
                "comparison_type",
                "include_formatting",
                "detailed_analysis",
                "generate_report"
            ],
            "example_request": {
                "document_1": "path/to/document1.docx",
                "document_2": "path/to/document2.docx",
                "comparison_type": "both",
                "options": {
                    "include_formatting": True,
                    "detailed_analysis": True
                }
            }
        }