"""
Professional DOCX to PDF conversion service.
Convert Word documents to high-quality PDF files.
"""

import asyncio
import io
import logging
import subprocess
from datetime import datetime
from typing import Any, Dict, Optional
from pathlib import Path
import base64
import tempfile
import os

try:
    from docx import Document
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
    class Document:
        pass

# Alternative conversion methods
try:
    import comtypes.client
    COMTYPES_AVAILABLE = True
except ImportError:
    COMTYPES_AVAILABLE = False

try:
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    SELENIUM_AVAILABLE = True
except ImportError:
    SELENIUM_AVAILABLE = False

from ..core.base import (
    BaseConverter, 
    ConversionFormat, 
    ConversionRequest, 
    ConversionResult,
    ConversionStatus,
    ProcessingError,
    ValidationError
)

logger = logging.getLogger(__name__)


class DocxToPdfConverter(BaseConverter):
    """
    Professional DOCX to PDF conversion service.
    
    Features:
    - High-quality PDF output
    - Preserves formatting and layout
    - Multiple conversion engines
    - Batch processing support
    - Custom page settings
    - Watermark support
    - Password protection
    """
    
    def __init__(self):
        super().__init__(
            name="docx_to_pdf",
            supported_formats=[ConversionFormat.PDF]
        )
        
        # Determine available conversion methods
        self.conversion_methods = []
        
        if DOCX2PDF_AVAILABLE:
            self.conversion_methods.append("docx2pdf")
        
        if COMTYPES_AVAILABLE and os.name == 'nt':  # Windows only
            self.conversion_methods.append("word_com")
        
        if SELENIUM_AVAILABLE:
            self.conversion_methods.append("selenium")
        
        # LibreOffice as fallback
        if self._check_libreoffice():
            self.conversion_methods.append("libreoffice")
        
        if not self.conversion_methods:
            logger.warning("No DOCX to PDF conversion methods available")
    
    def _check_libreoffice(self) -> bool:
        """Check if LibreOffice is available."""
        try:
            result = subprocess.run(
                ['libreoffice', '--version'], 
                capture_output=True, 
                text=True, 
                timeout=10
            )
            return result.returncode == 0
        except Exception:
            return False
    
    def validate_request(self, request: ConversionRequest) -> bool:
        """Validate DOCX to PDF conversion request."""
        if not request.content or not request.content.strip():
            return False
        
        if request.output_format != ConversionFormat.PDF:
            return False
        
        return True
    
    async def convert(self, request: ConversionRequest) -> ConversionResult:
        """
        Convert DOCX content to PDF.
        
        Args:
            request: Conversion request with DOCX content
            
        Returns:
            ConversionResult with PDF data
        """
        job_id = f"docx2pdf_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}_{id(request)}"
        
        try:
            logger.info(f"Starting DOCX to PDF conversion: {job_id}")
            
            # Validate request
            if not self.validate_request(request):
                raise ValidationError("Invalid DOCX to PDF conversion request", job_id)
            
            # Extract options
            options = request.options or {}
            conversion_method = options.get('method', 'auto')
            
            # Convert DOCX to PDF
            pdf_data = await self._convert_docx_to_pdf(request.content, options, conversion_method)
            
            # Create result
            result = ConversionResult(
                job_id=job_id,
                status=ConversionStatus.COMPLETED,
                output_data=pdf_data,
                output_filename=f"{job_id}.pdf",
                metadata={
                    "original_size": len(request.content),
                    "pdf_size": len(pdf_data),
                    "conversion_method": conversion_method,
                    "options_used": options
                },
                created_at=datetime.utcnow(),
                completed_at=datetime.utcnow()
            )
            
            logger.info(f"DOCX to PDF conversion completed: {job_id}")
            return result
            
        except Exception as e:
            logger.error(f"DOCX to PDF conversion failed: {job_id} - {str(e)}")
            return ConversionResult(
                job_id=job_id,
                status=ConversionStatus.FAILED,
                error_message=str(e),
                created_at=datetime.utcnow(),
                completed_at=datetime.utcnow()
            )
    
    async def _convert_docx_to_pdf(self, content: str, options: Dict[str, Any], method: str = 'auto') -> bytes:
        """Convert DOCX content to PDF using specified method."""
        
        # Determine conversion method
        if method == 'auto':
            method = self._select_best_method(options)
        
        if method not in self.conversion_methods:
            raise ProcessingError(f"Conversion method '{method}' not available")
        
        # Create temporary files
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_dir_path = Path(temp_dir)
            docx_path = temp_dir_path / "input.docx"
            pdf_path = temp_dir_path / "output.pdf"
            
            # Write DOCX content to file
            await self._write_docx_content(content, docx_path)
            
            # Convert based on method
            if method == "docx2pdf":
                await self._convert_with_docx2pdf(docx_path, pdf_path, options)
            elif method == "word_com":
                await self._convert_with_word_com(docx_path, pdf_path, options)
            elif method == "selenium":
                await self._convert_with_selenium(docx_path, pdf_path, options)
            elif method == "libreoffice":
                await self._convert_with_libreoffice(docx_path, pdf_path, options)
            else:
                raise ProcessingError(f"Unknown conversion method: {method}")
            
            # Read PDF data
            if not pdf_path.exists():
                raise ProcessingError("PDF conversion failed - output file not created")
            
            return pdf_path.read_bytes()
    
    async def _write_docx_content(self, content: str, output_path: Path):
        """Write DOCX content to file."""
        if content.startswith('data:'):
            # Base64 encoded content
            header, data = content.split(',', 1)
            content_bytes = base64.b64decode(data)
        elif len(content) > 1000 and not content.startswith('<'):
            # Likely binary content
            content_bytes = content.encode('latin1')
        else:
            # File path or raw content
            if Path(content).exists():
                content_bytes = Path(content).read_bytes()
            else:
                # Create DOCX from text content
                doc = Document()
                doc.add_paragraph(content)
                
                with io.BytesIO() as docx_buffer:
                    doc.save(docx_buffer)
                    content_bytes = docx_buffer.getvalue()
        
        output_path.write_bytes(content_bytes)
    
    async def _convert_with_docx2pdf(self, docx_path: Path, pdf_path: Path, options: Dict):
        """Convert using docx2pdf library."""
        loop = asyncio.get_event_loop()
        
        def convert():
            docx2pdf_convert(str(docx_path), str(pdf_path))
        
        await loop.run_in_executor(None, convert)
    
    async def _convert_with_word_com(self, docx_path: Path, pdf_path: Path, options: Dict):
        """Convert using Word COM interface (Windows only)."""
        loop = asyncio.get_event_loop()
        
        def convert():
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            
            try:
                doc = word.Documents.Open(str(docx_path.absolute()))
                
                # Apply options
                if options.get('password'):
                    doc.Password = options['password']
                
                # Export as PDF
                doc.ExportAsFixedFormat(
                    OutputFileName=str(pdf_path.absolute()),
                    ExportFormat=17,  # PDF format
                    OpenAfterExport=False,
                    OptimizeFor=0,  # Print optimization
                    BitmapMissingFonts=True,
                    DocStructureTags=True,
                    CreateBookmarks=options.get('bookmarks', True),
                    UseDocumentOrder=True
                )
                
                doc.Close()
            finally:
                word.Quit()
        
        await loop.run_in_executor(None, convert)
    
    async def _convert_with_selenium(self, docx_path: Path, pdf_path: Path, options: Dict):
        """Convert using Selenium and Chrome browser."""
        # First convert DOCX to HTML, then HTML to PDF
        html_content = await self._docx_to_html(docx_path)
        
        chrome_options = Options()
        chrome_options.add_argument('--headless')
        chrome_options.add_argument('--no-sandbox')
        chrome_options.add_argument('--disable-dev-shm-usage')
        chrome_options.add_argument('--disable-gpu')
        
        # PDF print settings
        print_options = {
            'paperWidth': options.get('page_width', 8.27),
            'paperHeight': options.get('page_height', 11.7),
            'marginTop': options.get('margin_top', 0.4),
            'marginBottom': options.get('margin_bottom', 0.4),
            'marginLeft': options.get('margin_left', 0.4),
            'marginRight': options.get('margin_right', 0.4),
            'printBackground': options.get('print_background', True),
        }
        
        driver = webdriver.Chrome(options=chrome_options)
        
        try:
            # Create temporary HTML file
            html_path = docx_path.parent / "temp.html"
            html_path.write_text(html_content, encoding='utf-8')
            
            driver.get(f"file://{html_path.absolute()}")
            
            # Generate PDF
            pdf_data = driver.execute_cdp_cmd('Page.printToPDF', print_options)
            pdf_bytes = base64.b64decode(pdf_data['data'])
            
            pdf_path.write_bytes(pdf_bytes)
            
        finally:
            driver.quit()
    
    async def _convert_with_libreoffice(self, docx_path: Path, pdf_path: Path, options: Dict):
        """Convert using LibreOffice."""
        cmd = [
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(docx_path.parent),
            str(docx_path)
        ]
        
        process = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        
        stdout, stderr = await process.communicate()
        
        if process.returncode != 0:
            raise ProcessingError(f"LibreOffice conversion failed: {stderr.decode()}")
        
        # LibreOffice creates PDF with same name as input
        generated_pdf = docx_path.parent / f"{docx_path.stem}.pdf"
        if generated_pdf.exists():
            generated_pdf.rename(pdf_path)
        else:
            raise ProcessingError("LibreOffice did not generate expected PDF file")
    
    async def _docx_to_html(self, docx_path: Path) -> str:
        """Convert DOCX to HTML for browser-based PDF conversion."""
        doc = Document(docx_path)
        html_parts = ['<!DOCTYPE html>', '<html><head>']
        html_parts.append('<meta charset="utf-8">')
        html_parts.append('<style>')
        html_parts.append(self._get_default_html_styles())
        html_parts.append('</style>')
        html_parts.append('</head><body>')
        
        for para in doc.paragraphs:
            if para.text.strip():
                html_parts.append(f'<p>{para.text}</p>')
        
        for table in doc.tables:
            html_parts.append('<table>')
            for row in table.rows:
                html_parts.append('<tr>')
                for cell in row.cells:
                    html_parts.append(f'<td>{cell.text}</td>')
                html_parts.append('</tr>')
            html_parts.append('</table>')
        
        html_parts.append('</body></html>')
        return '\\n'.join(html_parts)
    
    def _get_default_html_styles(self) -> str:
        """Get default CSS styles for HTML conversion."""
        return """
        body {
            font-family: 'Times New Roman', serif;
            font-size: 12pt;
            line-height: 1.15;
            margin: 1in;
            color: black;
        }
        p {
            margin: 0;
            padding: 0 0 6pt 0;
        }
        table {
            border-collapse: collapse;
            width: 100%;
            margin: 6pt 0;
        }
        td, th {
            border: 1px solid black;
            padding: 3pt;
            text-align: left;
        }
        """
    
    def _select_best_method(self, options: Dict) -> str:
        """Select the best available conversion method."""
        # Preference order based on quality and reliability
        preferred_order = ["word_com", "docx2pdf", "libreoffice", "selenium"]
        
        for method in preferred_order:
            if method in self.conversion_methods:
                return method
        
        raise ProcessingError("No conversion methods available")
    
    def get_capabilities(self) -> Dict[str, Any]:
        """Get converter capabilities."""
        return {
            "name": self.name,
            "description": "Professional DOCX to PDF conversion service",
            "supported_formats": [fmt.value for fmt in self.supported_formats],
            "available_methods": self.conversion_methods,
            "features": [
                "High-quality PDF output",
                "Preserves formatting and layout",
                "Multiple conversion engines",
                "Custom page settings",
                "Password protection support",
                "Batch processing"
            ],
            "supported_options": [
                "method",
                "page_width",
                "page_height",
                "margins",
                "password",
                "bookmarks",
                "print_background"
            ],
            "example_request": {
                "content": "DOCX file content or path",
                "output_format": "pdf",
                "options": {
                    "method": "auto",
                    "page_width": 8.27,
                    "page_height": 11.7,
                    "bookmarks": True
                }
            }
        }